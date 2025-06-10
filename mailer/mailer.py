# import smtplib
# from email.mime.multipart import MIMEMultipart
# from email.mime.text import MIMEText
# import google.generativeai as genai
# import os
# from dotenv import load_dotenv
# load_dotenv()

# from scrapper import scrape_dawn_articles

# # ==== CONFIGURATION ====
# myEmail = os.environ.get("EMAIL_ADDRESS")
# myPassword = os.environ.get("EMAIL_PASSWORD")
# gemini_api_key = os.environ.get("GEMINI_API_KEY")

# prompt = "Explain the basics of attention mechanism in transformers."

# # ==== GENERATE TEXT FROM GEMINI ====
# genai.configure(api_key=gemini_api_key)

# model = genai.GenerativeModel("gemini-2.0-flash")
# response = model.generate_content(prompt)

# gemini_text = response.text.strip()

# # ==== EMAIL SETUP ====
# smtp_server = "smtp.gmail.com"
# smtp_port = 587

# msg = MIMEMultipart()
# msg['Subject'] = 'Gemini API Response'
# msg['From'] = myEmail
# msg['To'] = "demo.abdullah.dev@gmail.com"

# body = f"Prompt:\n{prompt}\n\nGemini's Response:\n{gemini_text}"
# msg.attach(MIMEText(body, 'plain'))

# # ==== SEND EMAIL ====
# server = smtplib.SMTP(smtp_server, smtp_port)
# server.starttls()
# server.login(myEmail, myPassword)
# server.sendmail(myEmail, msg['To'], msg.as_string())
# server.quit()

# print("Email sent successfully.")

# ---------------------------------------------------------------------------------

# import smtplib
# from email.mime.multipart import MIMEMultipart
# from email.mime.text import MIMEText
# import google.generativeai as genai
# import os
# from dotenv import load_dotenv
# from scrapper import scrape_dawn_articles
# from datetime import datetime

# load_dotenv()

# # ==== CONFIGURATION ====
# myEmail = os.environ.get("EMAIL_ADDRESS")
# myPassword = os.environ.get("EMAIL_PASSWORD")
# gemini_api_key = os.environ.get("GEMINI_API_KEY")

# # ==== SETUP GEMINI ====
# genai.configure(api_key=gemini_api_key)
# model = genai.GenerativeModel("gemini-2.0-flash")

# # ==== SCRAPE ARTICLES ====
# articles = scrape_dawn_articles()  # Uses today’s date by default

# # ==== GENERATE PROMPT + GEMINI RESPONSES ====
# email_body = ""

# for article in articles:
#     title = article["title"]
#     content = article["content"]

#     prompt = f"""
# Read this article and provide two distinct lists based on its content, strictly following the format below:

# 1. Words: List all good vocabulary words from the article, along with a concise meaning for each. Format each entry as "word: meaning", and place each complete entry on a new line. Do not use any bullet points, dashes, or other special formatting. 
# 2. Phrases and Idioms: List different phrases and idioms found in the article, along with a concise meaning for each. Format each entry as "phrase: meaning", and place each complete entry on a new line. Do not use any bullet points, dashes, or other special formatting. 

# Title: {title}
# Content: {content}
# """.strip()

#     try:
#         response = model.generate_content(prompt)
#         gemini_text = response.text.strip()
#     except Exception as e:
#         gemini_text = f"Error generating response: {e}"

#     # email_body += f"\n\n=== Article: {title} ===\n\nPrompt:\n{prompt}\n\nGemini's Response:\n{gemini_text}\n"
#     email_body += f"\n\n=== Article: {title} ===\n\n{gemini_text}\n"

# # ==== EMAIL SETUP ====
# msg = MIMEMultipart()
# msg['Subject'] = f"Gemini Vocabulary Report - {datetime.today().strftime('%Y-%m-%d')}"
# msg['From'] = myEmail
# msg['To'] = "demo.abdullah.dev@gmail.com"

# msg.attach(MIMEText(email_body, 'plain'))

# # ==== SEND EMAIL ====
# smtp_server = "smtp.gmail.com"
# smtp_port = 587

# server = smtplib.SMTP(smtp_server, smtp_port)
# server.starttls()
# server.login(myEmail, myPassword)
# server.sendmail(myEmail, msg['To'], msg.as_string())
# server.quit()

# print("Email sent successfully.")










# import os
# import smtplib
# from email.mime.multipart import MIMEMultipart
# from email.mime.application import MIMEApplication
# from email.mime.text import MIMEText
# from datetime import datetime
# from pylatex import Document, Section, Subsection, Tabular
# import google.generativeai as genai
# from dotenv import load_dotenv
# from scrapper import scrape_dawn_articles  # Your existing scraper

# # ==== LOAD CONFIG ====
# load_dotenv()
# myEmail = os.environ.get("EMAIL_ADDRESS")
# myPassword = os.environ.get("EMAIL_PASSWORD")
# gemini_api_key = os.environ.get("GEMINI_API_KEY")

# # ==== SETUP GEMINI ====
# genai.configure(api_key=gemini_api_key)
# model = genai.GenerativeModel("gemini-2.0-flash")

# # ==== PARSE GEMINI OUTPUT ====
# def parse_response(text):
#     parts = text.split("2. Phrases and Idioms:")
#     words_raw = parts[0].replace("1. Words:", "").strip()
#     phrases_raw = parts[1].strip() if len(parts) > 1 else ""
    
#     def parse_lines(raw):
#         lines = raw.splitlines()
#         entries = [line.strip().split(":", 1) for line in lines if ":" in line]
#         return [(w.strip(), m.strip()) for w, m in entries]

#     return parse_lines(words_raw), parse_lines(phrases_raw)

# # ==== GENERATE LATEX PDF ====
# def generate_latex_pdf(all_data, pdf_filename="vocab_summary"):
#     doc = Document()
#     for title, words, phrases in all_data:
#         with doc.create(Section(title)):
#             with doc.create(Subsection("Words")):
#                 with doc.create(Tabular("ll")) as table:
#                     table.add_row(("Word", "Meaning"))
#                     table.add_hline()
#                     for word, meaning in words:
#                         table.add_row(word, meaning)
#             with doc.create(Subsection("Phrases and Idioms")):
#                 with doc.create(Tabular("ll")) as table:
#                     table.add_row(("Phrase", "Meaning"))
#                     table.add_hline()
#                     for phrase, meaning in phrases:
#                         table.add_row(phrase, meaning)
#     # doc.generate_pdf(pdf_filename, clean_tex=False)
#     doc.generate_pdf(pdf_filename, clean_tex=False, compiler='pdflatex')

#     return f"{pdf_filename}.pdf"

# # ==== MAIN WORKFLOW ====
# articles = scrape_dawn_articles()
# all_article_data = []

# for article in articles:
#     title = article["title"]
#     content = article["content"]

#     prompt = f"""
# Read this article and provide two distinct lists based on its content, strictly following the format below:

# 1. Words: List all good vocabulary words from the article, along with a concise meaning for each. Format each entry as "word: meaning", and place each complete entry on a new line. Do not use any bullet points, dashes, or other special formatting. 
# 2. Phrases and Idioms: List different phrases and idioms found in the article, along with a concise meaning for each. Format each entry as "phrase: meaning", and place each complete entry on a new line. Do not use any bullet points, dashes, or other special formatting. 

# - Make sure the meaning is not more than 10 words.
# Title: {title}
# Content: {content}
# """.strip()

#     try:
#         response = model.generate_content(prompt)
#         gemini_text = response.text.strip()
#         words, phrases = parse_response(gemini_text)
#         all_article_data.append((title, words, phrases))
#     except Exception as e:
#         print(f"Error generating response for article '{title}': {e}")

# # ==== GENERATE PDF ====
# pdf_filename = f"gemini_vocab_{datetime.today().strftime('%Y-%m-%d')}"
# pdf_path = generate_latex_pdf(all_article_data, pdf_filename)

# # ==== EMAIL SETUP ====
# msg = MIMEMultipart()
# msg['Subject'] = f"Gemini Vocabulary PDF Report - {datetime.today().strftime('%Y-%m-%d')}"
# msg['From'] = myEmail
# msg['To'] = "demo.abdullah.dev@gmail.com"

# msg.attach(MIMEText("Please find attached the Gemini-generated vocabulary report in PDF format.", 'plain'))

# # Attach PDF
# with open(pdf_path, "rb") as f:
#     part = MIMEApplication(f.read(), _subtype="pdf")
#     part.add_header('Content-Disposition', 'attachment', filename=os.path.basename(pdf_path))
#     msg.attach(part)

# # ==== SEND EMAIL ====
# smtp_server = "smtp.gmail.com"
# smtp_port = 587

# server = smtplib.SMTP(smtp_server, smtp_port)
# server.starttls()
# server.login(myEmail, myPassword)
# server.sendmail(myEmail, msg['To'], msg.as_string())
# server.quit()

# print("Email with PDF sent successfully.")







import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
import google.generativeai as genai
import os
from dotenv import load_dotenv
from scrapper import scrape_dawn_articles
from datetime import datetime
from docx import Document
from io import BytesIO

load_dotenv()

# ==== CONFIGURATION ====
myEmail = os.environ.get("EMAIL_ADDRESS")
myPassword = os.environ.get("EMAIL_PASSWORD")
gemini_api_key = os.environ.get("GEMINI_API_KEY")

# ==== SETUP GEMINI ====
genai.configure(api_key=gemini_api_key)
model = genai.GenerativeModel("gemini-2.0-flash")

# ==== SCRAPE ARTICLES ====
articles = scrape_dawn_articles()  # Uses today’s date by default

# ==== GENERATE PROMPT + GEMINI RESPONSES ====
doc = Document()
doc.add_heading(f"Vocabulary Report - {datetime.today().strftime('%Y-%m-%d')}", 0)

for article in articles:
    title = article["title"]
    content = article["content"]

    prompt = f"""
Read this article and provide two distinct lists based on its content, strictly following the format below:

1. Words: List all good vocabulary words from the article, along with a concise meaning for each. Format each entry as "word: meaning", and place each complete entry on a new line. Do not use any bullet points, dashes, or other special formatting. 
2. Phrases and Idioms: List different phrases and idioms found in the article, along with a concise meaning for each. Format each entry as "phrase: meaning", and place each complete entry on a new line. Do not use any bullet points, dashes, or other special formatting. 

Title: {title}
Content: {content}
""".strip()

    try:
        response = model.generate_content(prompt)
        gemini_text = response.text.strip()
    except Exception as e:
        gemini_text = f"Error generating response: {e}"

    doc.add_heading(f"Article: {title}", level=1)
    doc.add_paragraph(gemini_text)

# Save document to a bytes buffer instead of disk
doc_buffer = BytesIO()
doc.save(doc_buffer)
doc_buffer.seek(0)

# ==== EMAIL SETUP ====
# msg = MIMEMultipart()
# msg['Subject'] = f"Gemini Vocabulary Report - {datetime.today().strftime('%Y-%m-%d')}"
# msg['From'] = myEmail
# msg['To'] = "demo.abdullah.dev@gmail.com"

# # Attach a plain text body for email clients that don't render attachments well
# msg.attach(MIMEText("Please find attached the Gemini vocabulary report in DOCX format.", 'plain'))

# # Attach the docx file
# part = MIMEBase('application', "vnd.openxmlformats-officedocument.wordprocessingml.document")
# part.set_payload(doc_buffer.read())
# encoders.encode_base64(part)
# part.add_header('Content-Disposition', 'attachment', filename=f"Gemini_Vocab_Report_{datetime.today().strftime('%Y-%m-%d')}.docx")
# msg.attach(part)

# # ==== SEND EMAIL ====
# smtp_server = "smtp.gmail.com"
# smtp_port = 587

# server = smtplib.SMTP(smtp_server, smtp_port)
# server.starttls()
# server.login(myEmail, myPassword)
# server.sendmail(myEmail, msg['To'], msg.as_string())
# server.quit()

# print("Email with DOCX attachment sent successfully.")
# ==== EMAIL SETUP ====
recipient_emails = ["demo.abdullah.dev@gmail.com", "ausama.bese22seecs@seecs.edu.pk", "hareemfatima2244h@gmail.com","2683pbs@pphs.edu.pk","younasadvocate3@gmail.com","Imranahmedlashari21@gmail.com"]

msg = MIMEMultipart()
msg['Subject'] = f"Vocabulary Report - {datetime.today().strftime('%Y-%m-%d')}"
msg['From'] = myEmail
msg['To'] = ", ".join(recipient_emails)  # For the email header

# Attach plain text body
msg.attach(MIMEText("Please find attached the Gemini vocabulary report in DOCX format.", 'plain'))

# Attach the DOCX file
doc_buffer.seek(0)  # Ensure buffer is at beginning
part = MIMEBase('application', "vnd.openxmlformats-officedocument.wordprocessingml.document")
part.set_payload(doc_buffer.read())
encoders.encode_base64(part)
part.add_header('Content-Disposition', 'attachment', filename=f"Vocab_Report_{datetime.today().strftime('%Y-%m-%d')}.docx")
msg.attach(part)

# ==== SEND EMAIL ====
smtp_server = "smtp.gmail.com"
smtp_port = 587

server = smtplib.SMTP(smtp_server, smtp_port)
server.starttls()
server.login(myEmail, myPassword)
server.sendmail(myEmail, recipient_emails, msg.as_string())  # Use list here
server.quit()

print("Email with DOCX attachment sent to multiple recipients successfully.")
